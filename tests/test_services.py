"""Tests for DocumentProcessingService."""

import tempfile
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock, patch

import pytest

from apps.documents.models import Document
from apps.documents.services import DocumentProcessingService


@pytest.mark.unit
class TestDocumentProcessingService:
    """Tests for DocumentProcessingService."""

    @pytest.fixture
    def service(self) -> DocumentProcessingService:
        """Create service instance."""
        return DocumentProcessingService()

    @pytest.fixture
    def txt_file(self) -> Any:
        """Create a temporary TXT file."""
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False, encoding="utf-8"
        ) as f:
            f.write("This is a test document with some invoice details.\n")
            f.write("Total amount: $100.00\n")
            f.write("Payment due: 30 days\n")
            return Path(f.name)

    @pytest.fixture
    def pdf_document(self, db: Any) -> MagicMock:
        """Create a mock PDF document."""
        doc = MagicMock(spec=Document)
        doc.file_type = Document.FileType.PDF
        doc.file = MagicMock()
        doc.file.path = "/tmp/test.pdf"
        return doc

    @pytest.fixture
    def txt_document(self, db: Any, txt_file: Path) -> MagicMock:
        """Create a mock TXT document."""
        doc = MagicMock(spec=Document)
        doc.file_type = Document.FileType.TXT
        doc.file = MagicMock()
        doc.file.path = str(txt_file)
        return doc

    @pytest.fixture
    def docx_document(self, db: Any) -> MagicMock:
        """Create a mock DOCX document."""
        doc = MagicMock(spec=Document)
        doc.file_type = Document.FileType.DOCX
        doc.file = MagicMock()
        doc.file.path = "/tmp/test.docx"
        return doc

    @pytest.fixture
    def image_document(self, db: Any) -> MagicMock:
        """Create a mock IMAGE document."""
        doc = MagicMock(spec=Document)
        doc.file_type = Document.FileType.IMAGE
        doc.file = MagicMock()
        doc.file.path = "/tmp/test.jpg"
        return doc

    def test_extract_text_from_txt(
        self, service: DocumentProcessingService, txt_document: MagicMock
    ) -> None:
        """Test extracting text from TXT file."""
        text = service.extract_text(txt_document)

        assert "test document" in text
        assert "invoice" in text
        assert "$100.00" in text

    @patch.object(DocumentProcessingService, "_extract_text_from_pdf")
    def test_extract_text_from_pdf(
        self,
        mock_extract: MagicMock,
        service: DocumentProcessingService,
        pdf_document: MagicMock,
    ) -> None:
        """Test extracting text from PDF file."""
        mock_extract.return_value = "PDF content extracted"

        text = service.extract_text(pdf_document)

        assert text == "PDF content extracted"
        mock_extract.assert_called_once_with("/tmp/test.pdf")

    @patch.object(DocumentProcessingService, "_extract_text_from_docx")
    def test_extract_text_from_docx(
        self,
        mock_extract: MagicMock,
        service: DocumentProcessingService,
        docx_document: MagicMock,
    ) -> None:
        """Test extracting text from DOCX file."""
        mock_extract.return_value = "DOCX content extracted"

        text = service.extract_text(docx_document)

        assert text == "DOCX content extracted"
        mock_extract.assert_called_once_with("/tmp/test.docx")

    @patch.object(DocumentProcessingService, "_extract_text_from_image")
    def test_extract_text_from_image(
        self,
        mock_extract: MagicMock,
        service: DocumentProcessingService,
        image_document: MagicMock,
    ) -> None:
        """Test extracting text from IMAGE file (OCR placeholder)."""
        mock_extract.return_value = ""

        text = service.extract_text(image_document)

        assert text == ""
        mock_extract.assert_called_once_with("/tmp/test.jpg")

    def test_extract_text_unsupported_file_type(
        self, service: DocumentProcessingService, db: Any
    ) -> None:
        """Тест извлечения текста из неподдерживаемого типа файла."""
        doc = MagicMock()
        doc.file_type = "UNSUPPORTED"
        doc.file = MagicMock()
        doc.file.path = "/tmp/test.xyz"

        with pytest.raises(ValueError) as exc_info:
            service.extract_text(doc)

        assert "Неподдерживаемый тип файла" in str(exc_info.value)

    def test_extract_text_file_error(
        self, service: DocumentProcessingService, db: Any
    ) -> None:
        """Test extracting text when file doesn't exist."""
        doc = MagicMock()
        doc.file_type = Document.FileType.TXT
        doc.file = MagicMock()
        doc.file.path = "/nonexistent/path.txt"

        with pytest.raises(FileNotFoundError):
            service.extract_text(doc)

    def test_classify_document_invoice(
        self, service: DocumentProcessingService
    ) -> None:
        """Test classifying document as invoice."""
        text = (
            "This is an invoice for services. Total amount due: $500. Payment required."
        )

        classification, confidence = service.classify_document(text)

        assert classification == "invoice"
        assert confidence > 0

    def test_classify_document_contract(
        self, service: DocumentProcessingService
    ) -> None:
        """Test classifying document as contract."""
        text = "This agreement outlines the terms and conditions. Signature required."

        classification, confidence = service.classify_document(text)

        assert classification == "contract"
        assert confidence > 0

    def test_classify_document_report(self, service: DocumentProcessingService) -> None:
        """Test classifying document as report."""
        text = (
            "Annual report with summary and analysis. Findings and conclusion included."
        )

        classification, confidence = service.classify_document(text)

        assert classification == "report"
        assert confidence > 0

    def test_classify_document_letter(self, service: DocumentProcessingService) -> None:
        """Test classifying document as letter."""
        text = "Dear Mr. Smith, sincerely yours. Best regards."

        classification, confidence = service.classify_document(text)

        assert classification == "letter"
        assert confidence > 0

    def test_classify_document_unclassified(
        self, service: DocumentProcessingService
    ) -> None:
        """Test classifying document with no matching keywords."""
        text = "Random text with no recognizable keywords whatsoever."

        classification, confidence = service.classify_document(text)

        assert classification == "unclassified"
        assert confidence == 0.0

    def test_classify_document_empty_text(
        self, service: DocumentProcessingService
    ) -> None:
        """Test classifying empty document."""
        classification, confidence = service.classify_document("")

        assert classification == "unclassified"
        assert confidence == 0.0

    def test_extract_text_from_image_returns_empty(
        self, service: DocumentProcessingService
    ) -> None:
        """Test that OCR placeholder returns empty string."""
        result = service._extract_text_from_image("/tmp/test.jpg")
        assert result == ""


@pytest.mark.unit
class TestDocumentProcessingServiceRealExtraction:
    """Tests for real file extraction methods."""

    @pytest.fixture
    def service(self) -> DocumentProcessingService:
        """Create service instance."""
        return DocumentProcessingService()

    def test_extract_text_from_docx_real_file(
        self, service: DocumentProcessingService
    ) -> None:
        """Test extracting text from a real DOCX file."""
        import docx

        with tempfile.NamedTemporaryFile(mode="wb", suffix=".docx", delete=False) as f:
            docx_path = f.name

        doc = docx.Document()
        doc.add_paragraph("Test DOCX Content")
        doc.add_paragraph("Second paragraph")
        doc.save(docx_path)

        try:
            result = service._extract_text_from_docx(docx_path)
            assert "Test DOCX Content" in result
            assert "Second paragraph" in result
        finally:
            Path(docx_path).unlink(missing_ok=True)

    @pytest.mark.filterwarnings("ignore::DeprecationWarning")
    def test_extract_text_from_pdf_real_file(
        self, service: DocumentProcessingService
    ) -> None:
        """Test extracting text from a real PDF file."""
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        with tempfile.NamedTemporaryFile(mode="wb", suffix=".pdf", delete=False) as f:
            pdf_path = f.name

        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.drawString(100, 750, "Test PDF Content")
        c.drawString(100, 700, "Second line of text")
        c.save()

        try:
            result = service._extract_text_from_pdf(pdf_path)
            assert "Test PDF Content" in result
            assert "Second line" in result
        finally:
            Path(pdf_path).unlink(missing_ok=True)
